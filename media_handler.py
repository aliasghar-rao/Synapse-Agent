"""
Media Handler module for Drive-Manager Pro.
Provides advanced media file operations for images, videos, audio, and documents.
"""

import os
import sys
import time
import threading
import hashlib
import json
import tempfile
from datetime import datetime

# Optional imports - will be used if available
try:
    import cv2
    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False

try:
    from PIL import Image, ExifTags
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import moviepy.editor as mp
    MOVIEPY_AVAILABLE = True
except ImportError:
    MOVIEPY_AVAILABLE = False

try:
    import mutagen
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False

try:
    import PyPDF2
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class MediaType:
    """Enumeration of media types"""
    IMAGE = "image"
    VIDEO = "video"
    AUDIO = "audio"
    DOCUMENT = "document"
    UNKNOWN = "unknown"

class MediaMetadata:
    """Class to store and manage media file metadata"""
    
    def __init__(self, file_path):
        self.file_path = file_path
        self.filename = os.path.basename(file_path)
        self.extension = os.path.splitext(file_path)[1].lower()
        self.size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
        self.media_type = self._determine_media_type()
        self.last_modified = os.path.getmtime(file_path) if os.path.exists(file_path) else 0
        
        # Basic metadata common to all types
        self.metadata = {
            'filename': self.filename,
            'path': file_path,
            'size': self.size,
            'formatted_size': self._format_size(self.size),
            'extension': self.extension,
            'media_type': self.media_type,
            'last_modified': datetime.fromtimestamp(self.last_modified).strftime('%Y-%m-%d %H:%M:%S'),
            'hash': None
        }
        
        # Type-specific metadata will be added by the analyze method
        
    def _determine_media_type(self):
        """Determine media type based on file extension"""
        # Image extensions
        if self.extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.svg']:
            return MediaType.IMAGE
        # Video extensions
        elif self.extension in ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm', '.mpeg', '.3gp']:
            return MediaType.VIDEO
        # Audio extensions
        elif self.extension in ['.mp3', '.wav', '.ogg', '.flac', '.aac', '.m4a', '.wma']:
            return MediaType.AUDIO
        # Document extensions
        elif self.extension in ['.pdf', '.doc', '.docx', '.txt', '.rtf', '.odt', '.ppt', '.pptx', '.xls', '.xlsx']:
            return MediaType.DOCUMENT
        # Unknown
        else:
            return MediaType.UNKNOWN
    
    def _format_size(self, size_bytes):
        """Format size in bytes to human-readable format"""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024 or unit == 'TB':
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024
    
    def calculate_hash(self):
        """Calculate SHA-256 hash of file"""
        if not os.path.exists(self.file_path):
            return None
            
        try:
            hash_sha256 = hashlib.sha256()
            with open(self.file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            hash_value = hash_sha256.hexdigest()
            self.metadata['hash'] = hash_value
            return hash_value
        except Exception as e:
            print(f"Error calculating hash: {e}")
            return None
    
    def analyze(self):
        """Analyze media file to extract detailed metadata"""
        try:
            if self.media_type == MediaType.IMAGE:
                self._analyze_image()
            elif self.media_type == MediaType.VIDEO:
                self._analyze_video()
            elif self.media_type == MediaType.AUDIO:
                self._analyze_audio()
            elif self.media_type == MediaType.DOCUMENT:
                self._analyze_document()
                
            return self.metadata
        except Exception as e:
            print(f"Error analyzing {self.file_path}: {e}")
            return self.metadata
    
    def _analyze_image(self):
        """Extract metadata from image files"""
        if not PIL_AVAILABLE:
            self.metadata['status'] = 'PIL not available'
            return
            
        try:
            img = Image.open(self.file_path)
            
            # Basic image information
            self.metadata['width'] = img.width
            self.metadata['height'] = img.height
            self.metadata['aspect_ratio'] = round(img.width / img.height, 2)
            self.metadata['format'] = img.format
            self.metadata['mode'] = img.mode
            self.metadata['frames'] = getattr(img, 'n_frames', 1)
            self.metadata['is_animated'] = getattr(img, 'is_animated', False)
            
            # EXIF data (if available)
            exif_data = {}
            if hasattr(img, '_getexif') and img._getexif():
                exif = {
                    ExifTags.TAGS[k]: v
                    for k, v in img._getexif().items()
                    if k in ExifTags.TAGS and isinstance(v, (str, int, float, bytes))
                }
                
                # Extract commonly useful EXIF data
                if 'DateTimeOriginal' in exif:
                    exif_data['date_taken'] = exif['DateTimeOriginal']
                if 'Make' in exif:
                    exif_data['camera_make'] = exif['Make']
                if 'Model' in exif:
                    exif_data['camera_model'] = exif['Model']
                if 'FocalLength' in exif:
                    exif_data['focal_length'] = str(exif['FocalLength'])
                if 'ExposureTime' in exif:
                    exif_data['exposure_time'] = str(exif['ExposureTime'])
                if 'ISOSpeedRatings' in exif:
                    exif_data['iso'] = exif['ISOSpeedRatings']
                    
                self.metadata['exif'] = exif_data
            
            # Additional analysis with OpenCV if available
            if CV2_AVAILABLE:
                img_cv = cv2.imread(self.file_path)
                if img_cv is not None:
                    # Color analysis
                    average_color_per_row = cv2.mean(img_cv)[0:3]
                    self.metadata['average_color'] = {
                        'b': int(average_color_per_row[0]),
                        'g': int(average_color_per_row[1]),
                        'r': int(average_color_per_row[2]),
                        'hex': '#{:02x}{:02x}{:02x}'.format(
                            int(average_color_per_row[2]), 
                            int(average_color_per_row[1]), 
                            int(average_color_per_row[0])
                        )
                    }
                    
                    # Face detection
                    face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
                    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                    faces = face_cascade.detectMultiScale(gray, 1.3, 5)
                    self.metadata['faces_detected'] = len(faces)
        
        except Exception as e:
            self.metadata['error'] = str(e)
    
    def _analyze_video(self):
        """Extract metadata from video files"""
        if not CV2_AVAILABLE:
            if MOVIEPY_AVAILABLE:
                try:
                    clip = mp.VideoFileClip(self.file_path)
                    self.metadata['duration'] = clip.duration
                    self.metadata['fps'] = clip.fps
                    self.metadata['width'] = clip.size[0]
                    self.metadata['height'] = clip.size[1]
                    self.metadata['aspect_ratio'] = round(clip.size[0] / clip.size[1], 2)
                    self.metadata['frames'] = int(clip.fps * clip.duration)
                    self.metadata['audio_present'] = clip.audio is not None
                    clip.close()
                except Exception as e:
                    self.metadata['error'] = str(e)
            else:
                self.metadata['status'] = 'OpenCV and MoviePy not available'
            return
            
        try:
            cap = cv2.VideoCapture(self.file_path)
            
            if not cap.isOpened():
                self.metadata['error'] = 'Could not open video file'
                return
                
            # Basic video information
            self.metadata['fps'] = cap.get(cv2.CAP_PROP_FPS)
            self.metadata['width'] = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            self.metadata['height'] = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            self.metadata['aspect_ratio'] = round(self.metadata['width'] / self.metadata['height'], 2)
            self.metadata['frames'] = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            self.metadata['duration'] = self.metadata['frames'] / self.metadata['fps']
            self.metadata['formatted_duration'] = self._format_duration(self.metadata['duration'])
            
            # Extract first frame for thumbnail
            ret, frame = cap.read()
            if ret:
                # Analyze first frame for color
                average_color_per_row = cv2.mean(frame)[0:3]
                self.metadata['average_color'] = {
                    'b': int(average_color_per_row[0]),
                    'g': int(average_color_per_row[1]),
                    'r': int(average_color_per_row[2]),
                    'hex': '#{:02x}{:02x}{:02x}'.format(
                        int(average_color_per_row[2]), 
                        int(average_color_per_row[1]), 
                        int(average_color_per_row[0])
                    )
                }
                
                # Face detection on first frame
                face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                faces = face_cascade.detectMultiScale(gray, 1.3, 5)
                self.metadata['faces_detected_first_frame'] = len(faces)
                
            cap.release()
            
            # Get audio information with MoviePy if available
            if MOVIEPY_AVAILABLE:
                try:
                    clip = mp.VideoFileClip(self.file_path)
                    self.metadata['audio_present'] = clip.audio is not None
                    if clip.audio:
                        self.metadata['audio_fps'] = clip.audio.fps
                        self.metadata['audio_duration'] = clip.audio.duration
                    clip.close()
                except Exception:
                    self.metadata['audio_present'] = "Unknown"
        
        except Exception as e:
            self.metadata['error'] = str(e)
    
    def _analyze_audio(self):
        """Extract metadata from audio files"""
        if not MUTAGEN_AVAILABLE:
            if MOVIEPY_AVAILABLE:
                try:
                    audio = mp.AudioFileClip(self.file_path)
                    self.metadata['duration'] = audio.duration
                    self.metadata['fps'] = audio.fps
                    self.metadata['formatted_duration'] = self._format_duration(audio.duration)
                    audio.close()
                except Exception as e:
                    self.metadata['error'] = str(e)
            else:
                self.metadata['status'] = 'Mutagen and MoviePy not available'
            return
            
        try:
            # Using mutagen for most audio formats
            audio = mutagen.File(self.file_path)
            
            if audio is None:
                self.metadata['error'] = 'Could not parse audio file'
                return
                
            # Basic audio information
            if hasattr(audio, 'info'):
                if hasattr(audio.info, 'length'):
                    self.metadata['duration'] = audio.info.length
                    self.metadata['formatted_duration'] = self._format_duration(audio.info.length)
                if hasattr(audio.info, 'bitrate'):
                    self.metadata['bitrate'] = audio.info.bitrate
                if hasattr(audio.info, 'sample_rate'):
                    self.metadata['sample_rate'] = audio.info.sample_rate
                if hasattr(audio.info, 'channels'):
                    self.metadata['channels'] = audio.info.channels
            
            # Tags and metadata
            tags = {}
            if isinstance(audio, mutagen.id3.ID3FileType):
                # MP3 and similar formats with ID3 tags
                for key in audio:
                    if key.startswith('TPE1'):  # Artist
                        tags['artist'] = str(audio[key])
                    elif key.startswith('TIT2'):  # Title
                        tags['title'] = str(audio[key])
                    elif key.startswith('TALB'):  # Album
                        tags['album'] = str(audio[key])
                    elif key.startswith('TDRC'):  # Year
                        tags['year'] = str(audio[key])
                    elif key.startswith('TCON'):  # Genre
                        tags['genre'] = str(audio[key])
            else:
                # Other formats
                for key in audio:
                    if isinstance(audio[key], list) and len(audio[key]) > 0:
                        tags[key] = str(audio[key][0])
                    else:
                        tags[key] = str(audio[key])
            
            self.metadata['tags'] = tags
        
        except Exception as e:
            self.metadata['error'] = str(e)
    
    def _analyze_document(self):
        """Extract metadata from document files"""
        if self.extension == '.pdf':
            if not PYPDF_AVAILABLE:
                self.metadata['status'] = 'PyPDF2 not available'
                return
                
            try:
                with open(self.file_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    
                    # Basic PDF information
                    self.metadata['pages'] = len(pdf.pages)
                    
                    # Document info (title, author, etc.)
                    if pdf.metadata:
                        doc_info = {}
                        for key, value in pdf.metadata.items():
                            if key and value and isinstance(value, (str, int, float)):
                                doc_info[key] = value
                        self.metadata['document_info'] = doc_info
                    
                    # Extract text from first page for preview
                    if len(pdf.pages) > 0:
                        first_page = pdf.pages[0]
                        text = first_page.extract_text()
                        if text:
                            self.metadata['first_page_text'] = text[:500] + '...' if len(text) > 500 else text
                    
            except Exception as e:
                self.metadata['error'] = str(e)
                
        elif self.extension == '.docx':
            if not DOCX_AVAILABLE:
                self.metadata['status'] = 'python-docx not available'
                return
                
            try:
                doc = docx.Document(self.file_path)
                
                # Basic document information
                self.metadata['paragraphs'] = len(doc.paragraphs)
                self.metadata['sections'] = len(doc.sections)
                
                # Count tables and images
                tables_count = len(doc.tables)
                self.metadata['tables'] = tables_count
                
                # Document properties
                core_properties = doc.core_properties
                props = {}
                if core_properties.author:
                    props['author'] = core_properties.author
                if core_properties.title:
                    props['title'] = core_properties.title
                if core_properties.comments:
                    props['comments'] = core_properties.comments
                if core_properties.created:
                    props['created'] = core_properties.created.isoformat()
                if core_properties.modified:
                    props['modified'] = core_properties.modified.isoformat()
                
                self.metadata['document_properties'] = props
                
                # Extract some text for preview
                text = ""
                for i, para in enumerate(doc.paragraphs):
                    if i < 5:  # First 5 paragraphs
                        text += para.text + "\n"
                    else:
                        break
                        
                if text:
                    self.metadata['preview_text'] = text[:500] + '...' if len(text) > 500 else text
                
            except Exception as e:
                self.metadata['error'] = str(e)
                
        elif self.extension == '.txt':
            try:
                with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read(1000)  # Read first 1000 characters
                    lines = text.count('\n') + 1
                    words = len(text.split())
                    
                    self.metadata['preview'] = text[:500] + '...' if len(text) > 500 else text
                    self.metadata['lines_sample'] = lines
                    self.metadata['words_sample'] = words
                    
            except Exception as e:
                self.metadata['error'] = str(e)
        
        # Other document types would be analyzed here
    
    def _format_duration(self, seconds):
        """Format duration in seconds to HH:MM:SS format"""
        hours, remainder = divmod(seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        return f"{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d}"


class MediaHandler:
    """Manager for media file operations"""
    
    def __init__(self, cache_dir=None):
        self.cache_dir = cache_dir or os.path.join(tempfile.gettempdir(), 'drivemanager_media_cache')
        os.makedirs(self.cache_dir, exist_ok=True)
        self.metadata_cache = {}
        
    def get_metadata(self, file_path, force_refresh=False):
        """Get metadata for a media file"""
        if not os.path.exists(file_path):
            return None
            
        # Check cache first
        if not force_refresh and file_path in self.metadata_cache:
            # Check if file has been modified since last analysis
            last_modified = os.path.getmtime(file_path)
            if last_modified <= self.metadata_cache[file_path].get('last_modified', 0):
                return self.metadata_cache[file_path]
        
        # Create media metadata object and analyze
        media = MediaMetadata(file_path)
        metadata = media.analyze()
        
        # Cache the result
        self.metadata_cache[file_path] = metadata
        
        return metadata
    
    def batch_analyze(self, file_paths, callback=None):
        """Analyze multiple files and return their metadata"""
        results = {}
        
        for i, file_path in enumerate(file_paths):
            metadata = self.get_metadata(file_path)
            results[file_path] = metadata
            
            if callback:
                progress = (i + 1) / len(file_paths) * 100
                callback(progress, file_path, metadata)
        
        return results
    
    def generate_thumbnail(self, file_path, width=200, height=200):
        """Generate a thumbnail for a media file"""
        if not os.path.exists(file_path):
            return None
            
        media_type = self._get_media_type(file_path)
        
        # Generate thumbnail based on media type
        if media_type == MediaType.IMAGE and PIL_AVAILABLE:
            return self._generate_image_thumbnail(file_path, width, height)
        elif media_type == MediaType.VIDEO and CV2_AVAILABLE:
            return self._generate_video_thumbnail(file_path, width, height)
        elif media_type == MediaType.DOCUMENT and PYPDF_AVAILABLE:
            return self._generate_document_thumbnail(file_path, width, height)
        else:
            # Generate a generic thumbnail
            return self._generate_generic_thumbnail(file_path, width, height)
    
    def _get_media_type(self, file_path):
        """Get media type for a file"""
        extension = os.path.splitext(file_path)[1].lower()
        
        # Image extensions
        if extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.svg']:
            return MediaType.IMAGE
        # Video extensions
        elif extension in ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm', '.mpeg', '.3gp']:
            return MediaType.VIDEO
        # Audio extensions
        elif extension in ['.mp3', '.wav', '.ogg', '.flac', '.aac', '.m4a', '.wma']:
            return MediaType.AUDIO
        # Document extensions
        elif extension in ['.pdf', '.doc', '.docx', '.txt', '.rtf', '.odt', '.ppt', '.pptx', '.xls', '.xlsx']:
            return MediaType.DOCUMENT
        # Unknown
        else:
            return MediaType.UNKNOWN
    
    def _generate_image_thumbnail(self, file_path, width, height):
        """Generate thumbnail for an image file"""
        try:
            # Create a thumbnail using PIL
            img = Image.open(file_path)
            img.thumbnail((width, height))
            
            # Save to a temp file
            thumbnail_path = os.path.join(self.cache_dir, 
                                        f"thumb_{hashlib.md5(file_path.encode()).hexdigest()}.jpg")
            img.save(thumbnail_path, "JPEG")
            
            return thumbnail_path
        except Exception as e:
            print(f"Error generating image thumbnail: {e}")
            return None
    
    def _generate_video_thumbnail(self, file_path, width, height):
        """Generate thumbnail for a video file"""
        try:
            # Create a thumbnail using OpenCV
            cap = cv2.VideoCapture(file_path)
            
            if not cap.isOpened():
                return None
                
            # Try to get a frame from the 10% mark of the video
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            if frame_count > 0:
                cap.set(cv2.CAP_PROP_POS_FRAMES, min(10, frame_count // 10))
            
            ret, frame = cap.read()
            cap.release()
            
            if not ret:
                return None
                
            # Resize the frame
            frame = cv2.resize(frame, (width, height))
            
            # Save to a temp file
            thumbnail_path = os.path.join(self.cache_dir, 
                                        f"thumb_{hashlib.md5(file_path.encode()).hexdigest()}.jpg")
            cv2.imwrite(thumbnail_path, frame)
            
            return thumbnail_path
        except Exception as e:
            print(f"Error generating video thumbnail: {e}")
            return None
    
    def _generate_document_thumbnail(self, file_path, width, height):
        """Generate thumbnail for a document file"""
        try:
            extension = os.path.splitext(file_path)[1].lower()
            
            if extension == '.pdf' and PYPDF_AVAILABLE:
                # For PDF files, render the first page
                with open(file_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    
                    if len(pdf.pages) > 0 and PIL_AVAILABLE:
                        # This is a simplified approach; in a real implementation,
                        # you would render the PDF page as an image
                        img = Image.new('RGB', (width, height), color='white')
                        
                        # Save to a temp file
                        thumbnail_path = os.path.join(self.cache_dir, 
                                                    f"thumb_{hashlib.md5(file_path.encode()).hexdigest()}.jpg")
                        img.save(thumbnail_path, "JPEG")
                        
                        return thumbnail_path
            
            # For other document types or if PDF rendering fails, use a generic thumbnail
            return self._generate_generic_thumbnail(file_path, width, height)
            
        except Exception as e:
            print(f"Error generating document thumbnail: {e}")
            return self._generate_generic_thumbnail(file_path, width, height)
    
    def _generate_generic_thumbnail(self, file_path, width, height):
        """Generate a generic thumbnail based on file type"""
        try:
            if not PIL_AVAILABLE:
                return None
                
            # Create a blank image with file extension text
            img = Image.new('RGB', (width, height), color='white')
            
            # Save to a temp file
            thumbnail_path = os.path.join(self.cache_dir, 
                                        f"thumb_{hashlib.md5(file_path.encode()).hexdigest()}.jpg")
            img.save(thumbnail_path, "JPEG")
            
            return thumbnail_path
        except Exception as e:
            print(f"Error generating generic thumbnail: {e}")
            return None
    
    def extract_preview(self, file_path, max_size=1024*1024):
        """Extract a preview of the file contents"""
        if not os.path.exists(file_path):
            return None
            
        media_type = self._get_media_type(file_path)
        
        # Extract preview based on media type
        if media_type == MediaType.DOCUMENT:
            return self._extract_document_preview(file_path, max_size)
        elif media_type == MediaType.IMAGE and PIL_AVAILABLE:
            return self._generate_image_thumbnail(file_path, 800, 600)
        elif media_type == MediaType.VIDEO and CV2_AVAILABLE:
            return self._generate_video_thumbnail(file_path, 800, 600)
        else:
            return None
    
    def _extract_document_preview(self, file_path, max_size=1024*1024):
        """Extract a preview from a document file"""
        try:
            extension = os.path.splitext(file_path)[1].lower()
            
            if extension == '.pdf' and PYPDF_AVAILABLE:
                with open(file_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    
                    if len(pdf.pages) > 0:
                        # Extract text from first page
                        text = pdf.pages[0].extract_text()
                        return text[:2000] + '...' if len(text) > 2000 else text
            
            elif extension == '.docx' and DOCX_AVAILABLE:
                doc = docx.Document(file_path)
                
                # Extract some text for preview
                text = ""
                for i, para in enumerate(doc.paragraphs):
                    if i < 10:  # First 10 paragraphs
                        text += para.text + "\n"
                    else:
                        break
                        
                return text[:2000] + '...' if len(text) > 2000 else text
            
            elif extension == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read(2000)  # Read first 2000 characters
                    return content + '...' if len(content) >= 2000 else content
            
            # Other document types would be handled here
            return None
            
        except Exception as e:
            print(f"Error extracting document preview: {e}")
            return None
    
    def convert_media(self, source_path, target_format, output_path=None):
        """Convert a media file to a different format"""
        if not os.path.exists(source_path):
            return {'success': False, 'error': 'Source file does not exist'}
            
        # Determine source media type
        source_extension = os.path.splitext(source_path)[1].lower()
        media_type = self._get_media_type(source_path)
        
        # Create default output path if not provided
        if not output_path:
            base_name = os.path.splitext(os.path.basename(source_path))[0]
            output_path = os.path.join(os.path.dirname(source_path), f"{base_name}.{target_format}")
        
        # Handle different media types
        if media_type == MediaType.IMAGE:
            return self._convert_image(source_path, target_format, output_path)
        elif media_type == MediaType.VIDEO:
            return self._convert_video(source_path, target_format, output_path)
        elif media_type == MediaType.AUDIO:
            return self._convert_audio(source_path, target_format, output_path)
        else:
            return {'success': False, 'error': 'Unsupported media type for conversion'}
    
    def _convert_image(self, source_path, target_format, output_path):
        """Convert an image file to a different format"""
        if not PIL_AVAILABLE:
            return {'success': False, 'error': 'PIL not available for image conversion'}
            
        try:
            img = Image.open(source_path)
            
            # Handle formats that don't support alpha channel
            if img.mode == 'RGBA' and target_format.lower() in ['jpg', 'jpeg']:
                bg = Image.new('RGB', img.size, (255, 255, 255))
                bg.paste(img, mask=img.split()[3])  # 3 is the alpha channel
                img = bg
            
            img.save(output_path, target_format.upper())
            return {
                'success': True, 
                'output_path': output_path,
                'message': f"Converted image to {target_format}"
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _convert_video(self, source_path, target_format, output_path):
        """Convert a video file to a different format"""
        if not MOVIEPY_AVAILABLE:
            return {'success': False, 'error': 'MoviePy not available for video conversion'}
            
        try:
            video = mp.VideoFileClip(source_path)
            
            # Set codec based on target format
            codec = None
            if target_format.lower() == 'mp4':
                codec = 'libx264'
            elif target_format.lower() == 'webm':
                codec = 'libvpx'
            elif target_format.lower() == 'avi':
                codec = 'rawvideo'
            
            video.write_videofile(output_path, codec=codec)
            video.close()
            
            return {
                'success': True, 
                'output_path': output_path,
                'message': f"Converted video to {target_format}"
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _convert_audio(self, source_path, target_format, output_path):
        """Convert an audio file to a different format"""
        if not MOVIEPY_AVAILABLE:
            return {'success': False, 'error': 'MoviePy not available for audio conversion'}
            
        try:
            audio = mp.AudioFileClip(source_path)
            audio.write_audiofile(output_path)
            audio.close()
            
            return {
                'success': True, 
                'output_path': output_path,
                'message': f"Converted audio to {target_format}"
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def extract_frames(self, video_path, output_dir=None, frame_rate=1):
        """Extract frames from a video file at specified frame rate"""
        if not os.path.exists(video_path):
            return {'success': False, 'error': 'Video file does not exist'}
            
        if not CV2_AVAILABLE:
            return {'success': False, 'error': 'OpenCV not available for frame extraction'}
            
        # Create output directory if not provided
        if not output_dir:
            video_name = os.path.splitext(os.path.basename(video_path))[0]
            output_dir = os.path.join(os.path.dirname(video_path), f"{video_name}_frames")
        
        os.makedirs(output_dir, exist_ok=True)
        
        try:
            cap = cv2.VideoCapture(video_path)
            
            if not cap.isOpened():
                return {'success': False, 'error': 'Could not open video file'}
                
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            
            # Calculate frame interval based on requested frame rate
            if frame_rate >= fps:
                frame_interval = 1  # Extract every frame
            else:
                frame_interval = int(fps / frame_rate)
            
            extracted_frames = 0
            current_frame = 0
            
            while True:
                ret, frame = cap.read()
                
                if not ret:
                    break
                
                if current_frame % frame_interval == 0:
                    frame_path = os.path.join(output_dir, f"frame_{extracted_frames:05d}.jpg")
                    cv2.imwrite(frame_path, frame)
                    extracted_frames += 1
                
                current_frame += 1
                
            cap.release()
            
            return {
                'success': True,
                'output_dir': output_dir,
                'extracted_frames': extracted_frames,
                'message': f"Extracted {extracted_frames} frames to {output_dir}"
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def extract_audio(self, video_path, output_path=None):
        """Extract audio track from a video file"""
        if not os.path.exists(video_path):
            return {'success': False, 'error': 'Video file does not exist'}
            
        if not MOVIEPY_AVAILABLE:
            return {'success': False, 'error': 'MoviePy not available for audio extraction'}
            
        # Create output path if not provided
        if not output_path:
            video_name = os.path.splitext(os.path.basename(video_path))[0]
            output_path = os.path.join(os.path.dirname(video_path), f"{video_name}.mp3")
        
        try:
            video = mp.VideoFileClip(video_path)
            
            if video.audio is None:
                video.close()
                return {'success': False, 'error': 'Video does not contain audio'}
                
            video.audio.write_audiofile(output_path)
            video.close()
            
            return {
                'success': True,
                'output_path': output_path,
                'message': f"Extracted audio to {output_path}"
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def resize_image(self, image_path, width, height, output_path=None, maintain_aspect_ratio=True):
        """Resize an image to specified dimensions"""
        if not os.path.exists(image_path):
            return {'success': False, 'error': 'Image file does not exist'}
            
        if not PIL_AVAILABLE:
            return {'success': False, 'error': 'PIL not available for image resizing'}
            
        # Create output path if not provided
        if not output_path:
            base_name, extension = os.path.splitext(image_path)
            output_path = f"{base_name}_resized{extension}"
        
        try:
            img = Image.open(image_path)
            
            if maintain_aspect_ratio:
                # Calculate new dimensions while maintaining aspect ratio
                img_width, img_height = img.size
                aspect_ratio = img_width / img_height
                
                if width / height > aspect_ratio:
                    # Width is the limiting factor
                    new_width = int(height * aspect_ratio)
                    new_height = height
                else:
                    # Height is the limiting factor
                    new_width = width
                    new_height = int(width / aspect_ratio)
                    
                resized_img = img.resize((new_width, new_height), Image.LANCZOS)
            else:
                # Resize without maintaining aspect ratio
                resized_img = img.resize((width, height), Image.LANCZOS)
            
            resized_img.save(output_path)
            
            return {
                'success': True,
                'output_path': output_path,
                'width': resized_img.width,
                'height': resized_img.height,
                'message': f"Resized image to {resized_img.width}x{resized_img.height}"
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def create_video_from_images(self, image_paths, output_path, fps=30, resolution=None):
        """Create a video from a sequence of images"""
        if not image_paths:
            return {'success': False, 'error': 'No images provided'}
            
        if not CV2_AVAILABLE:
            return {'success': False, 'error': 'OpenCV not available for video creation'}
            
        try:
            # Read the first image to get dimensions
            img = cv2.imread(image_paths[0])
            if img is None:
                return {'success': False, 'error': f'Could not read image: {image_paths[0]}'}
                
            height, width, channels = img.shape
            
            # If resolution is specified, use it instead
            if resolution:
                width, height = resolution
            
            # Create video writer
            fourcc = cv2.VideoWriter_fourcc(*'mp4v')  # Use mp4v codec for MP4 format
            out = cv2.VideoWriter(output_path, fourcc, fps, (width, height))
            
            # Process each image
            for image_path in image_paths:
                img = cv2.imread(image_path)
                
                if img is None:
                    continue
                    
                # Resize if necessary
                if img.shape[1] != width or img.shape[0] != height:
                    img = cv2.resize(img, (width, height))
                    
                out.write(img)
                
            out.release()
            
            return {
                'success': True,
                'output_path': output_path,
                'frames': len(image_paths),
                'width': width,
                'height': height,
                'fps': fps,
                'message': f"Created video from {len(image_paths)} images"
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def analyze_folder(self, folder_path, callback=None):
        """Analyze all media files in a folder"""
        if not os.path.isdir(folder_path):
            return {'success': False, 'error': 'Folder does not exist'}
            
        try:
            media_files = []
            
            # Find all media files in the folder
            for root, dirs, files in os.walk(folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    media_type = self._get_media_type(file_path)
                    
                    if media_type != MediaType.UNKNOWN:
                        media_files.append(file_path)
            
            # Analyze each file
            results = self.batch_analyze(media_files, callback)
            
            # Compile statistics
            stats = {
                'total_files': len(media_files),
                'by_type': {
                    MediaType.IMAGE: 0,
                    MediaType.VIDEO: 0,
                    MediaType.AUDIO: 0,
                    MediaType.DOCUMENT: 0,
                    MediaType.UNKNOWN: 0
                },
                'total_size': 0
            }
            
            for file_path, metadata in results.items():
                media_type = metadata.get('media_type', MediaType.UNKNOWN)
                stats['by_type'][media_type] = stats['by_type'].get(media_type, 0) + 1
                stats['total_size'] += metadata.get('size', 0)
            
            return {
                'success': True,
                'stats': stats,
                'results': results
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}


# Create an instance for use in the application
media_handler = MediaHandler()